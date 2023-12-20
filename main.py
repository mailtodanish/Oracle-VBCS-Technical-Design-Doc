import os
import pathlib
import json

import docx
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Cm


class VBCS:
    def __init__(self, base_dir: str, project_name: str):
        self.base_dir = base_dir
        self.project_name = project_name
        self.mdoc = docx.Document()
        style = self.mdoc.styles["Normal"]
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.alignment = 1
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(10)

        pass

    def project_dir_exist(
        self,
    ) -> bool:
        if not pathlib.Path(self.base_dir).is_dir():
            print("Please create a project folder.")
            return False

        dirs = os.listdir(self.base_dir)
        if len(dirs) > 0:
            self.project_path = f"{self.base_dir}/{self.project_name}"
            return True

        return False

    def read_json_from_file(self, path):
        f = open(path)
        obj = json.load(f)
        f.close()
        return obj

    def basic(self) -> bool:
        path = f"{self.project_path}/visual-application.json"
        if not pathlib.Path(path).is_file():
            print(f"File is not present: {path}.")
            return False

        detail = self.read_json_from_file(path)
        self.mdoc.add_heading(
            f"TDD for {detail.get('rootURL')} VBCS App", 0
        ).alignment = 1
        self.mdoc.add_paragraph(
            f"The {detail.get('rootURL')} app is being developed in the VBCS environment, and its technical details are as follows:"
        )

        self.mdoc.add_heading(f"Application", 0).alignment = 0

        table = self.mdoc.add_table(rows=1, cols=4)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        hdr_cells = table.rows[0].cells

        cell1 = hdr_cells[0].paragraphs[0].add_run("ROOT URL")
        cell1.bold = True
        cell2 = hdr_cells[1].paragraphs[0].add_run("VERSION")
        cell2.bold = True
        cell3 = hdr_cells[2].paragraphs[0].add_run("SOURCE VERSION")
        cell3.bold = True
        cell4 = hdr_cells[3].paragraphs[0].add_run("VBCS VERSION")
        cell4.bold = True

        row = table.add_row()
        row.cells[0].merge(hdr_cells[0])
        row.cells[1].merge(hdr_cells[1])
        row.cells[2].merge(hdr_cells[2])
        row.cells[3].merge(hdr_cells[3])

        row = table.add_row()
        row.cells[0].merge(hdr_cells[0])
        row.cells[1].merge(hdr_cells[1])
        row.cells[2].merge(hdr_cells[2])
        row.cells[3].merge(hdr_cells[3])

        row = table.add_row().cells
        row[0].text = detail.get("rootURL", "")
        row[1].text = detail.get("version", "")
        row[2].text = detail.get("source.version", "")
        row[3].text = detail.get("vbcs.dt.version", "")

        table.style = "Table Grid"
        table.style.paragraph_format.alignment = 2
        table.style.font.color.rgb = RGBColor(54, 0, 0)

        return True

    def businessObjectFields(self, data):
        self.mdoc.add_heading("Fields", 2).alignment = 0

        table = self.mdoc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        cell1 = hdr_cells[0].paragraphs[0].add_run("TYPE")
        cell1.bold = True
        cell2 = hdr_cells[1].paragraphs[0].add_run("NAME")
        cell2.bold = True
        cell3 = hdr_cells[2].paragraphs[0].add_run("REQUIRED")
        cell3.bold = True
        cell4 = hdr_cells[3].paragraphs[0].add_run("UNIQUE")
        cell4.bold = True
        cell5 = hdr_cells[4].paragraphs[0].add_run("DISPLAY LABEL")
        cell5.bold = True
        row = table.add_row()
        row.cells[0].merge(hdr_cells[0])
        row.cells[1].merge(hdr_cells[1])
        row.cells[2].merge(hdr_cells[2])
        row.cells[3].merge(hdr_cells[3])
        row.cells[4].merge(hdr_cells[4])

        for k in data:
            row = table.add_row().cells
            row[0].text = str(k["type"])
            row[1].text = str(k.get("name"))
            row[2].text = str(k.get("required"))
            row[3].text = str(k.get("unique"))
            row[4].text = str(k["displayLabel"])

        table.style = "Table Grid"
        table.style.paragraph_format.alignment = 2
        table.style.font.color.rgb = RGBColor(54, 0, 0)

    def businessObjectDetails(self, path):
        # parse entity JSON
        path = f"{path}/entity.json"
        if not pathlib.Path(path).is_file():
            print(f"Dir is not present: {path}")
            return False
        entity_detail = self.read_json_from_file(path)

        # table description
        self.mdoc.add_paragraph(entity_detail.get("description")).alignment=0

        table = self.mdoc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        cell1 = hdr_cells[0].paragraphs[0].add_run("PROPERTY")
        cell1.bold = True
        cell2 = hdr_cells[1].paragraphs[0].add_run("VALUE")
        cell2.bold = True

        row = table.add_row()
        row.cells[0].merge(hdr_cells[0])
        row.cells[1].merge(hdr_cells[1])

        for k in entity_detail.keys():
            if k in ["name", "displayLabel", "setupData"]:
                row = table.add_row().cells
                row[0].text = str(k)
                row[1].text = str(entity_detail[k])

        table.style = "Table Grid"
        table.style.paragraph_format.alignment = 2
        table.style.font.color.rgb = RGBColor(54, 0, 0)

        self.businessObjectFields(entity_detail["fields"])

        return True

    def buinessObjects(
        self,
    ) -> bool:
        objects_path = f"{self.project_path}/businessObjects/default/objects"
        if not pathlib.Path(objects_path).is_dir():
            print(f"Dir is not present: {objects_path}.")
            return False
        self.mdoc.add_paragraph("")
        self.mdoc.add_heading("Business Objects", 0).alignment = 0

        count = 1
        for dir in os.listdir(objects_path):
            self.mdoc.add_heading(f"{count}. {dir}", 1).alignment = 0
            self.businessObjectDetails(f"{objects_path}/{dir}")
            count = count + 1

        return True

    def serviceServers(self, data):
        self.mdoc.add_heading("Servers/Profile", 2).alignment = 0

        count = 1
        for s in data:
            vb = s.get("x-vb", {})
            self.mdoc.add_heading(
                f'{count}. {vb.get("profiles", "default")}', 3
            ).alignment = 0
            self.mdoc.add_paragraph(s.get("description")).alignment=0
            table = self.mdoc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            cell1 = hdr_cells[0].paragraphs[0].add_run("URL")
            cell1.bold = True
            cell2 = hdr_cells[1].paragraphs[0].add_run("AUTHENTICATION")
            cell2.bold = True
            cell3 = hdr_cells[2].paragraphs[0].add_run("ANONYMOUS ACCESS")
            cell3.bold = True

            row = table.add_row()
            row.cells[0].merge(hdr_cells[0])
            row.cells[1].merge(hdr_cells[1])
            row.cells[2].merge(hdr_cells[2])

            row = table.add_row().cells
            row[0].text = str(s["url"])
            auth = vb.get("authentication", {})
            row[1].text = str(auth.get("authenticated", {}).get("type", ""))
            row[2].text = str(vb.get("anonymousAccess", ""))
            table.style = "Table Grid"
            table.style.paragraph_format.alignment = 2
            table.style.font.color.rgb = RGBColor(54, 0, 0)
            count = count + 1

    def serviceDetails(self, path):
        # parse entity JSON
        path = f"{path}/openapi3.json"

        if not pathlib.Path(path).is_file():
            print(f"Dir is not present: {path}")
            return False

        entity_detail = self.read_json_from_file(path)

        self.mdoc.add_paragraph(entity_detail.get("info").get("description")).alignment=0

        table = self.mdoc.add_table(rows=1, cols=2)
        row = table.rows[0].cells
        row[0].text = "PROPERTY"
        row[1].text = "VALUE"

        for k in entity_detail.keys():
            if k == "info":
                row = table.add_row().cells
                row[0].text = "title"
                row[1].text = str(entity_detail[k]["title"])
            elif k == "servers":
                self.serviceServers(entity_detail[k])

        table.style = "Table Grid"
        table.style.paragraph_format.alignment = 2
        table.style.font.color.rgb = RGBColor(54, 0, 0)

    def services(self) -> bool:
        objects_path = f"{self.project_path}/services"
        if not pathlib.Path(objects_path).is_dir():
            print(f"Dir is not present: {objects_path}.")
            return False

        self.mdoc.add_paragraph("")
        self.mdoc.add_heading("Services/Connections", 0).alignment = 0

        count = 1
        for dir in os.listdir(objects_path):
            if not pathlib.Path(f"{objects_path}/{dir}").is_file():
                self.mdoc.add_heading(f"{count}. {dir}", 1).alignment = 0
                self.serviceDetails(f"{objects_path}/{dir}")
                count = count + 1

        return True

    def profiles(self) -> bool:
        path = f"{self.project_path}/settings"
        if not pathlib.Path(path).is_dir():
            print(f"Dir is not present: {path}.")
            return False

        self.mdoc.add_paragraph("")
        self.mdoc.add_heading("Profiles", 0).alignment = 0
        self.mdoc.add_paragraph("Below profiles has been created in applictaion.").alignment=0
        path = f"{path}/deployment-profiles.json"

        if not pathlib.Path(path).is_file():
            print(f"Dir is not present: {path}")
            return False

        profiles = self.read_json_from_file(path)

        table = self.mdoc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells

        cell1 = hdr_cells[0].paragraphs[0].add_run("ID")
        cell1.bold = True
        cell2 = hdr_cells[1].paragraphs[0].add_run("DISPLAY NAME")
        cell2.bold = True
        cell3 = hdr_cells[2].paragraphs[0].add_run("DESCRIPTION")
        cell3.bold = True

        row = table.add_row()
        row.cells[0].merge(hdr_cells[0])
        row.cells[1].merge(hdr_cells[1])
        row.cells[2].merge(hdr_cells[2])

        for p in profiles.get("profiles", []):
            row = table.add_row().cells
            row[0].text = p.get("id", "")
            row[1].text = p.get("displayName", "")
            row[2].text = p.get("description", "")

        table.style = "Table Grid"
        table.style.paragraph_format.alignment = 2
        table.style.font.color.rgb = RGBColor(54, 0, 0)

        return True

    def userRoles(self) -> bool:
        path = f"{self.project_path}/settings"
        if not pathlib.Path(path).is_dir():
            print(f"Dir is not present: {path}.")
            return False

        self.mdoc.add_paragraph("")
        self.mdoc.add_heading("User Roles", 0).alignment = 0
        self.mdoc.add_paragraph("Below roles has been created in applictaion.").alignment=0
        path = f"{path}/user-roles.json"

        if not pathlib.Path(path).is_file():
            print(f"Dir is not present: {path}")
            return False

        profiles = self.read_json_from_file(path)

        table = self.mdoc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        cell1 = hdr_cells[0].paragraphs[0].add_run("NAME")
        cell1.bold = True
        cell2 = hdr_cells[1].paragraphs[0].add_run("DESCRIPTION")
        cell2.bold = True
        cell3 = hdr_cells[2].paragraphs[0].add_run("ID")
        cell3.bold = True

        row = table.add_row()
        row.cells[0].merge(hdr_cells[0])
        row.cells[1].merge(hdr_cells[1])
        row.cells[2].merge(hdr_cells[2])

        for p in profiles.get("userroles", []):
            row = table.add_row().cells
            row[0].text = p.get("name", "")
            row[1].text = p.get("description", "")
            row[2].text = p.get("id", "")

        table.style = "Table Grid"
        table.style.paragraph_format.alignment = 2
        table.style.font.color.rgb = RGBColor(54, 0, 0)

        return True

    def actions(self, path):
        if not pathlib.Path(path).is_dir():
            print(f"File is not present: {path}.")
            return False

        for dir in os.listdir(path):
            if "json" not in dir:
                print(f"{path}/{dir}")
                continue
            obj = self.read_json_from_file(f"{path}/{dir}")
            self.mdoc.add_heading(f"Action : {dir}", 1).alignment = 0
            table = self.mdoc.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            cell1 = hdr_cells[0].paragraphs[0].add_run("NAME")
            cell1.bold = True
            cell2 = hdr_cells[1].paragraphs[0].add_run("DESCRIPTION")
            cell2.bold = True
            cell3 = hdr_cells[2].paragraphs[0].add_run("ROOT")
            cell3.bold = True

            row = table.add_row()
            row.cells[0].merge(hdr_cells[0])
            row.cells[1].merge(hdr_cells[1])
            row.cells[2].merge(hdr_cells[2])

            row = table.add_row().cells
            row[0].text = dir.replace(".json", "")
            row[1].text = obj.get("description", "")
            row[2].text = obj.get("root", "")
            table.style = "Table Grid"
            # Action Variables
            row = table.add_row().cells
            row[0].text = "VARIABLES"
            table.cell(3, 0).merge(table.cell(3, 2)).text = "Variables"
            table.cell(3, 0).merge(table.cell(3, 2)).paragraphs[0].bold = True
            vars = obj.get("variables", {})
            table = self.mdoc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            cell1 = hdr_cells[0].paragraphs[0].add_run("NAME")
            cell1.bold = True
            cell2 = hdr_cells[1].paragraphs[0].add_run("TYPE")
            cell2.bold = True
            cell3 = hdr_cells[2].paragraphs[0].add_run("REQUIRED")
            cell3.bold = True
            cell4 = hdr_cells[3].paragraphs[0].add_run("INPUT")
            cell4.bold = True

            row = table.add_row()
            row.cells[0].merge(hdr_cells[0])
            row.cells[1].merge(hdr_cells[1])
            row.cells[2].merge(hdr_cells[2])
            row.cells[3].merge(hdr_cells[3])

            for v in vars.keys():
                row = table.add_row().cells
                row[0].text = str(v)
                row[1].text = vars[v].get("type", "")
                row[2].text = str(vars[v].get("required", ""))
                row[3].text = vars[v].get("input", "")
            table.style = "Table Grid"
            table.style.paragraph_format.alignment = 2
            table.style.font.color.rgb = RGBColor(54, 0, 0)
        pass

    def app_flow_json(self, path):
        if not pathlib.Path(path).is_file():
            print(f"File is not present: {path}.")
            return False

        obj = self.read_json_from_file(path)

        # types on page
        types = obj.get("types", {})
        if types:
            self.mdoc.add_heading("Types", 1).alignment = 0
            table = self.mdoc.add_table(rows=1, cols=2)

            hdr_cells = table.rows[0].cells
            cell1 = hdr_cells[0].paragraphs[0].add_run("NAME")
            cell1.bold = True
            cell2 = hdr_cells[1].paragraphs[0].add_run("DESCRIPTION")
            cell2.bold = True

            row = table.add_row()
            row.cells[0].merge(hdr_cells[0])
            row.cells[1].merge(hdr_cells[1])

            for t in types.keys():
                row = table.add_row().cells
                row[0].text = str(t)
                row[1].text = str(types[t])
            table.style = "Table Grid"
            table.style.paragraph_format.alignment = 2
            table.style.font.color.rgb = RGBColor(54, 0, 0)

        # variables
        vars = obj.get("variables", {})
        if vars:
            self.mdoc.add_heading("Variables", 1).alignment = 0
            table = self.mdoc.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            cell1 = hdr_cells[0].paragraphs[0].add_run("NAME")
            cell1.bold = True
            cell2 = hdr_cells[1].paragraphs[0].add_run("TYPE")
            cell2.bold = True
            cell3 = hdr_cells[2].paragraphs[0].add_run("DEFAULT VALUE")
            cell3.bold = True

            row = table.add_row()
            row.cells[0].merge(hdr_cells[0])
            row.cells[1].merge(hdr_cells[1])
            row.cells[2].merge(hdr_cells[2])

            for t in vars.keys():
                row = table.add_row().cells
                row[0].text = str(t)
                row[1].text = str(
                    vars[t]["type"] if isinstance(vars[t]["type"], str) else "complex"
                )
                row[2].text = str(vars[t].get("defaultValue", ""))
            table.style = "Table Grid"
            table.style.paragraph_format.alignment = 2
            table.style.font.color.rgb = RGBColor(54, 0, 0)

        # security
        secu = obj.get("security", {})
        if secu:
            self.mdoc.add_heading("Security", 1).alignment = 0
            table = self.mdoc.add_table(rows=1, cols=2)

            hdr_cells = table.rows[0].cells
            cell1 = hdr_cells[0].paragraphs[0].add_run("PROPERTY")
            cell1.bold = True
            cell2 = hdr_cells[1].paragraphs[0].add_run("VALUE")
            cell2.bold = True

            row = table.add_row()
            row.cells[0].merge(hdr_cells[0])
            row.cells[1].merge(hdr_cells[1])

            access = secu.get("access", {})
            for a in access.keys():
                row = table.add_row().cells
                row[0].text = str(a)
                row[1].text = str(access.get(a))
            table.style = "Table Grid"
            table.style.paragraph_format.alignment = 2
            table.style.font.color.rgb = RGBColor(54, 0, 0)

        # eventListeners
        eventListeners = obj.get("eventListeners", None)
        if eventListeners:
            self.mdoc.add_heading("EventListeners", 1).alignment = 0
            table = self.mdoc.add_table(rows=1, cols=2)

            hdr_cells = table.rows[0].cells
            cell1 = hdr_cells[0].paragraphs[0].add_run("PROPERTY")
            cell1.bold = True
            cell2 = hdr_cells[1].paragraphs[0].add_run("VALUE")
            cell2.bold = True

            row = table.add_row()
            row.cells[0].merge(hdr_cells[0])
            row.cells[1].merge(hdr_cells[1])

            for a in eventListeners.keys():
                row = table.add_row().cells
                row[0].text = str(a)
                row[1].text = str(eventListeners.get(a))
            table.style = "Table Grid"
            table.style.paragraph_format.alignment = 2
            table.style.font.color.rgb = RGBColor(54, 0, 0)

    def webApps(self) -> bool:
        path = f"{self.project_path}/webApps"
        if not pathlib.Path(path).is_dir():
            print(f"File is not present: {path}.")
            return False

        self.mdoc.add_paragraph("")
        self.mdoc.add_heading("Web Apps", 0).alignment = 0

        count = 1
        for dir in os.listdir(path):
            if not pathlib.Path(path).is_dir():
                continue

            # issue in MAC
            if "DS_Store" in dir:
                continue

            self.mdoc.add_heading(f"{count}. {dir}", 1).alignment = 0
            # actions
            self.actions(f"{self.project_path}/webApps/{dir}/chains")
            # variables - app-flow.json
            self.app_flow_json(f"{self.project_path}/webApps/{dir}/app-flow.json")
            count = count + 1

            # pages
            self.pages(f"{self.project_path}/webApps/{dir}/flows/main/pages")

            # flow
            self.flow(f"{self.project_path}/webApps/{dir}/flows/main/flows")

    def pages(self, path):
        if not pathlib.Path(path).is_dir():
            print(f"File is not present: {path}.")
            return False
        self.mdoc.add_heading("Pages", 1).alignment = 0

        for dir in os.listdir(path):
            if "json" in dir:
                obj = self.read_json_from_file(f"{path}/{dir}")
                self.mdoc.add_heading(
                    f'Page:{obj.get("title") or obj.get("id")}', 2
                ).alignment = 0
                self.app_flow_json(f"{path}/{dir}")
            if "chain" in dir:
                self.mdoc.add_heading("Actions", 2).alignment = 0
                self.actions(f"{path}/{dir}")
        pass

    def flow(self, path):
        if not pathlib.Path(path).is_dir():
            print(f"File is not present: {path}.")
            return False
        self.mdoc.add_heading("Flow", 1).alignment = 0
        count = 1
        for dir in os.listdir(path):
            if not pathlib.Path(f"{path}/{dir}").is_dir():
                continue
            self.mdoc.add_heading(f"{count}. {dir}", 1).alignment = 0
            self.pages(f"{path}/{dir}/pages")
            count = count + 1

        pass

    def close_Doc(self) -> bool:
        self.mdoc.save(f"{pathlib.Path(__file__).parent}/TDD.docx")
        return True


if __name__ == "__main__":
    base_dir = f"{pathlib.Path(__file__).parent}/project"
    for dir in os.listdir(base_dir):
        # issue in MAC
        if "DS_Store" in dir:
            continue
        if pathlib.Path(f'{base_dir}/{dir}').is_dir():
            vb = VBCS(base_dir, dir)
            if vb.project_dir_exist():
                vb.basic()
                vb.buinessObjects()
                vb.services()
                vb.profiles()
                vb.userRoles()
                vb.webApps()
                vb.close_Doc()
            break

    
